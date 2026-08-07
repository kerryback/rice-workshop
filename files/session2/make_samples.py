#!/usr/bin/env python3
"""Generate the sample files used in the Academic Studio demo.

One small dataset, told nine ways, so a demo can move between file types
without changing subject. Run from files/session2/.
"""
import json
import os
import subprocess

import pandas as pd
from docx import Document
from docx.shared import Pt

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_files")
os.makedirs(OUT, exist_ok=True)

DATA = pd.DataFrame({
    "quarter": ["2024 Q1", "2024 Q2", "2024 Q3", "2024 Q4",
                "2025 Q1", "2025 Q2", "2025 Q3", "2025 Q4"],
    "revenue": [4.12, 4.38, 4.05, 5.21, 4.66, 4.91, 4.74, 5.88],
    "cost": [2.80, 2.95, 2.88, 3.40, 3.11, 3.24, 3.19, 3.72],
})
DATA["margin"] = ((DATA.revenue - DATA.cost) / DATA.revenue * 100).round(1)

MEMO_TITLE = "Quarterly Results, 2024-2025"
MEMO_BODY = [
    "Revenue grew from $4.12M in 2024 Q1 to $5.88M in 2025 Q4, an increase of "
    "43% over eight quarters.",
    "Margin has been steady between 28% and 37%. The fourth quarter is the "
    "strongest in both years, which is worth keeping in mind when comparing "
    "any single quarter against the one before it.",
    "The figures here are invented for teaching purposes.",
]


def csv_and_excel():
    DATA.to_csv(f"{OUT}/quarterly_revenue.csv", index=False)
    with pd.ExcelWriter(f"{OUT}/quarterly_revenue.xlsx", engine="openpyxl") as xl:
        DATA.to_excel(xl, sheet_name="Revenue", index=False)
        ws = xl.sheets["Revenue"]
        for col, width in zip("ABCD", (12, 12, 12, 12)):
            ws.column_dimensions[col].width = width
        for cell in ws[1]:
            cell.font = cell.font.copy(bold=True)


def python_script():
    with open(f"{OUT}/revenue_summary.py", "w") as f:
        f.write('''"""Summarise the quarterly revenue data."""
import pandas as pd

df = pd.read_csv("quarterly_revenue.csv")
df["margin"] = (df.revenue - df.cost) / df.revenue * 100

print(df.to_string(index=False))
print()
print(f"total revenue   {df.revenue.sum():.2f}M")
print(f"average margin  {df.margin.mean():.1f}%")
print(f"best quarter    {df.loc[df.revenue.idxmax(), 'quarter']}")
''')


def notebook():
    """A small notebook. json.dump does the escaping, so write plain strings."""
    md = ["# Quarterly revenue\n", "\n",
          "A first look at the numbers in `quarterly_revenue.csv`."]
    code1 = ["import pandas as pd\n", "\n",
             'df = pd.read_csv("quarterly_revenue.csv")\n', "df.head()"]
    code2 = ['df["margin"] = (df.revenue - df.cost) / df.revenue * 100\n',
             'df.plot(x="quarter", y="revenue", kind="bar", legend=False)']
    cell = lambda kind, src: {
        "cell_type": kind, "metadata": {}, "source": src,
        **({"execution_count": None, "outputs": []} if kind == "code" else {})}
    nb = {
        "cells": [cell("markdown", md), cell("code", code1), cell("code", code2)],
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python",
                           "name": "python3"},
            "language_info": {"name": "python", "version": "3.13"},
        },
        "nbformat": 4, "nbformat_minor": 5,
    }
    with open(f"{OUT}/revenue_analysis.ipynb", "w") as f:
        json.dump(nb, f, indent=1)


def markdown():
    lines = [f"# {MEMO_TITLE}", ""]
    for p in MEMO_BODY:
        lines += [p, ""]
    lines += ["| Quarter | Revenue | Cost | Margin |", "|---|---|---|---|"]
    for r in DATA.itertuples():
        lines.append(f"| {r.quarter} | {r.revenue:.2f} | {r.cost:.2f} | {r.margin}% |")
    with open(f"{OUT}/memo.md", "w") as f:
        f.write("\n".join(lines) + "\n")


def word():
    doc = Document()
    doc.add_heading(MEMO_TITLE, level=1)
    for p in MEMO_BODY:
        doc.add_paragraph(p)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for cell, head in zip(t.rows[0].cells, ("Quarter", "Revenue", "Cost", "Margin")):
        cell.text = head
    for r in DATA.itertuples():
        cells = t.add_row().cells
        for cell, val in zip(cells, (r.quarter, f"{r.revenue:.2f}",
                                     f"{r.cost:.2f}", f"{r.margin}%")):
            cell.text = val
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.save(f"{OUT}/memo.docx")


def latex_and_pdf():
    rows = " \\\\\n".join(
        f"{r.quarter} & {r.revenue:.2f} & {r.cost:.2f} & {r.margin}\\%"
        for r in DATA.itertuples())
    body = "\n\n".join(MEMO_BODY)
    tex = f"""\\documentclass[11pt]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{booktabs}}
\\title{{{MEMO_TITLE}}}
\\author{{Rice Business}}
\\date{{}}

\\begin{{document}}
\\maketitle

{body}

\\begin{{center}}
\\begin{{tabular}}{{lrrr}}
\\toprule
Quarter & Revenue & Cost & Margin \\\\
\\midrule
{rows} \\\\
\\bottomrule
\\end{{tabular}}
\\end{{center}}

\\end{{document}}
"""
    with open(f"{OUT}/memo.tex", "w") as f:
        f.write(tex)

    pdflatex = os.path.expanduser(
        "~/Library/TinyTeX/bin/universal-darwin/pdflatex")
    if not os.path.exists(pdflatex):
        print("  pdflatex not found; memo.pdf not built")
        return
    subprocess.run([pdflatex, "-interaction=nonstopmode", "memo.tex"],
                   cwd=OUT, capture_output=True, timeout=120)
    for ext in ("aux", "log", "out"):
        p = f"{OUT}/memo.{ext}"
        if os.path.exists(p):
            os.remove(p)


def html():
    rows = "\n".join(
        f"    <tr><td>{r.quarter}</td><td>{r.revenue:.2f}</td>"
        f"<td>{r.cost:.2f}</td><td>{r.margin}%</td></tr>"
        for r in DATA.itertuples())
    paras = "\n".join(f"  <p>{p}</p>" for p in MEMO_BODY)
    with open(f"{OUT}/dashboard.html", "w") as f:
        f.write(f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>{MEMO_TITLE}</title>
  <style>
    body {{ font: 16px/1.6 -apple-system, "Segoe UI", sans-serif;
           max-width: 46rem; margin: 3rem auto; padding: 0 1rem; color: #222; }}
    h1 {{ font-size: 1.6rem; }}
    table {{ border-collapse: collapse; width: 100%; margin-top: 1.5rem; }}
    th, td {{ text-align: right; padding: .45rem .6rem;
              border-bottom: 1px solid #ddd; }}
    th:first-child, td:first-child {{ text-align: left; }}
  </style>
</head>
<body>
  <h1>{MEMO_TITLE}</h1>
{paras}
  <table>
    <tr><th>Quarter</th><th>Revenue</th><th>Cost</th><th>Margin</th></tr>
{rows}
  </table>
</body>
</html>
""")


if __name__ == "__main__":
    csv_and_excel()
    python_script()
    notebook()
    markdown()
    word()
    latex_and_pdf()
    html()
    for name in sorted(os.listdir(OUT)):
        size = os.path.getsize(os.path.join(OUT, name))
        print(f"  {name:26} {size:>7,} bytes")
